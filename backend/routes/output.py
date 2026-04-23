from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from database import get_db
from models import User, JobDescription, TailoredResume
from routes.auth import get_current_user
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
import zipfile
import tempfile

router = APIRouter(prefix="/output", tags=["output"])


def set_spacing(para, before=0, after=0, line=1.15):
    fmt = para.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def build_docx(tailored_content: dict, output_path: str):
    doc = Document()
    
    # Ensure Hyperlink style exists with blue color
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    styles = doc.styles
    try:
        hyperlink_style = styles["Hyperlink"]
    except KeyError:
        hyperlink_style = styles.add_style("Hyperlink", 2)  # 2 = character style
        hyperlink_style.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        hyperlink_style.font.underline = True
    # Count content to decide spacing
    total_items = (
        len(tailored_content.get("work_experience", [])) +
        len(tailored_content.get("projects", [])) +
        len(tailored_content.get("education", [])) +
        len(tailored_content.get("leadership", [])) +
        len(tailored_content.get("certifications", [])) +
        len(tailored_content.get("achievements", []))
    )

    if total_items <= 5:
        SECTION_BEFORE = 20
        BULLET_LINE = 1.4
        BODY_AFTER = 6
        NAME_SIZE = 20
    elif total_items <= 8:
        SECTION_BEFORE = 14
        BULLET_LINE = 1.2
        BODY_AFTER = 4
        NAME_SIZE = 18
    else:
        SECTION_BEFORE = 10
        BULLET_LINE = 1.1
        BODY_AFTER = 2
        NAME_SIZE = 18

    # Margins - clean for ATS
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    def add_section_header(text):
        para = doc.add_paragraph()
        run = para.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(70, 70, 70)
        set_spacing(para, before=SECTION_BEFORE, after=4)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '999999')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return para

    def add_bullet(text):
        para = doc.add_paragraph(style='List Bullet')
        run = para.add_run(text)
        run.font.size = Pt(10)
        set_spacing(para, before=0, after=0, line=BULLET_LINE)
        return para

    def add_normal(text, size=10, italic=False):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(size)
        run.italic = italic
        set_spacing(para, before=0, after=BODY_AFTER)
        return para

    # ── Name ──────────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(tailored_content.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(NAME_SIZE)
    set_spacing(name_para, after=1)

    # ── Contact info ───────────────────────────────────────────────────────────
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def add_hyperlink(para, text, url):
        part = para.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run_elem = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "18")
        rPr.append(sz)
        run_elem.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        run_elem.append(t)
        hyperlink.append(run_elem)
        para._p.append(hyperlink)

    def add_separator(para):
        run = para.add_run(" | ")
        run.font.size = Pt(9)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(contact_para, after=6)

    contact_fields = []

    email = tailored_content.get("email")
    if email:
        contact_fields.append(("email", email))

    phone = tailored_content.get("phone")
    if phone:
        contact_fields.append(("phone", phone))

    linkedin = tailored_content.get("linkedin")
    if linkedin:
        contact_fields.append(("linkedin", linkedin))

    github = tailored_content.get("github")
    if github:
        contact_fields.append(("github", github))

    portfolio = tailored_content.get("portfolio")
    if portfolio:
        contact_fields.append(("portfolio", portfolio))

    for i, (field, val) in enumerate(contact_fields):
        if i > 0:
            add_separator(contact_para)
        if field in ("linkedin", "github", "portfolio"):
            # Ensure URL has http prefix
            url = val if val.startswith("http") else f"https://{val}"
            label = "LinkedIn" if field == "linkedin" else "GitHub" if field == "github" else "Portfolio"
            add_hyperlink(contact_para, label, url)
        else:
            run = contact_para.add_run(val)
            run.font.size = Pt(9)

    # ── Education ─────────────────────────────────────────────────────────────
    if tailored_content.get("education"):
        add_section_header("Education")
        for edu in tailored_content["education"]:
            para = doc.add_paragraph()
            set_spacing(para, before=6, after=0)
            run = para.add_run(edu.get("institution", ""))
            run.bold = True
            run.font.size = Pt(10)
            end = edu.get("end_year") or "Present"
            para.add_run(f"  {edu.get('start_year', '')} - {end}").font.size = Pt(10)
            add_normal(f"{edu.get('degree', '')} in {edu.get('field', '')}", italic=True)
            if edu.get("gpa"):
                add_normal(f"GPA: {edu['gpa']}")

    # ── Experience (Work + Leadership merged, deduplicated) ────────────────────
    experience = []
    seen = set()

    for item in tailored_content.get("work_experience", []):
        org = item.get("company", "")
        role = item.get("role", "")
        key = (org.strip().lower(), role.strip().lower())
        if key not in seen:
            seen.add(key)
            experience.append(("work", item))

    for item in tailored_content.get("leadership", []):
        org = item.get("organization", "")
        role = item.get("role", "")
        key = (org.strip().lower(), role.strip().lower())
        if key not in seen:
            seen.add(key)
            experience.append(("leadership", item))

    if experience:
        add_section_header("Experience")
        for kind, item in experience:
            para = doc.add_paragraph()
            set_spacing(para, before=6, after=0)
            org = item.get("company") or item.get("organization", "")
            run = para.add_run(org)
            run.bold = True
            run.font.size = Pt(10)
            start = item.get("start_date", "")
            end = item.get("end_date") or "Present"
            if start:
                para.add_run(f"  {start} - {end}").font.size = Pt(10)
            add_normal(item.get("role", ""), italic=True)
            bullets = item.get("bullets") or []
            if not bullets and item.get("impact"):
                bullets = [item["impact"]]
            for bullet in bullets:
                add_bullet(bullet)

    # ── Projects ───────────────────────────────────────────────────────────────
    if tailored_content.get("projects"):
        add_section_header("Projects")
        for project in tailored_content["projects"]:
            para = doc.add_paragraph()
            set_spacing(para, before=6, after=0)

            name = project.get("name", "")
            link = project.get("link", "")

            # Project name as plain bold text
            run = para.add_run(name)
            run.bold = True
            run.font.size = Pt(10)

            # Tech stack
            if project.get("tech_stack"):
                tech_run = para.add_run(f"  |  {', '.join(project['tech_stack'])}")
                tech_run.font.size = Pt(9)
                tech_run.italic = True

            # (Link) as clickable hyperlink at the end
            if link and link.strip():
                url = link if link.startswith("http") else f"https://{link}"
                separator = para.add_run("  ")
                separator.font.size = Pt(9)
                # Add opening bracket
                para.add_run("(").font.size = Pt(9)
                # Add hyperlink
                part = para.part
                r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
                hyperlink = OxmlElement("w:hyperlink")
                hyperlink.set(qn("r:id"), r_id)
                run_elem = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                rStyle = OxmlElement("w:rStyle")
                rStyle.set(qn("w:val"), "Hyperlink")
                rPr.append(rStyle)
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), "18")
                rPr.append(sz)
                run_elem.append(rPr)
                t = OxmlElement("w:t")
                t.text = "Link"
                run_elem.append(t)
                hyperlink.append(run_elem)
                para._p.append(hyperlink)
                # Add closing bracket
                para.add_run(")").font.size = Pt(9)

            for bullet in project.get("bullets", []):
                add_bullet(bullet)

    # ── Skills ─────────────────────────────────────────────────────────────────
    if tailored_content.get("skills"):
        add_section_header("Skills")
        skills = tailored_content["skills"]
        if skills.get("technical"):
            para = doc.add_paragraph()
            run = para.add_run("Technical:  ")
            run.bold = True
            run.font.size = Pt(10)
            para.add_run(", ".join(skills["technical"])).font.size = Pt(10)
        if skills.get("soft"):
            para = doc.add_paragraph()
            run = para.add_run("Soft:  ")
            run.bold = True
            run.font.size = Pt(10)
            para.add_run(", ".join(skills["soft"])).font.size = Pt(10)

    # ── Certifications ─────────────────────────────────────────────────────────
    if tailored_content.get("certifications"):
        add_section_header("Certifications")
        for cert in tailored_content["certifications"]:
            add_normal(f"{cert.get('name', '')} — {cert.get('issuer', '')} ({cert.get('year', '')})")

    # ── Achievements ───────────────────────────────────────────────────────────
    if tailored_content.get("achievements"):
        add_section_header("Achievements")
        for ach in tailored_content["achievements"]:
            add_bullet(f"{ach.get('title', '')} — {ach.get('description', '')}")

    # ── Volunteer ──────────────────────────────────────────────────────────────
    if tailored_content.get("volunteer"):
        add_section_header("Volunteer Experience")
        for vol in tailored_content["volunteer"]:
            para = doc.add_paragraph()
            set_spacing(para, before=6, after=0)
            run = para.add_run(vol.get("role", ""))
            run.bold = True
            run.font.size = Pt(10)
            add_normal(vol.get("organization", ""), italic=True)
            if vol.get("impact"):
                add_bullet(vol["impact"])

    doc.save(output_path)
   
         


@router.get("/download/{tailored_resume_id}")
def download_single(tailored_resume_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    record = db.query(TailoredResume).filter(
        TailoredResume.id == tailored_resume_id,
        TailoredResume.user_id == current_user.id
    ).first()
    if not record:
        raise HTTPException(status_code=404, detail="Tailored resume not found")

    jd = db.query(JobDescription).filter(JobDescription.id == record.jd_id).first()
    title = jd.title if jd and jd.title else f"resume_{tailored_resume_id}"
    safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()

    tailored_content = json.loads(record.content)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    build_docx(tailored_content, tmp.name)

    return FileResponse(
        path=tmp.name,
        filename=f"{safe_title}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@router.get("/download-all/{session_id}")
def download_all(session_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    jds = db.query(JobDescription).filter(
        JobDescription.session_id == session_id,
        JobDescription.user_id == current_user.id
    ).all()

    if not jds:
        raise HTTPException(status_code=404, detail="No JDs found in this session")

    tmp_dir = tempfile.mkdtemp()
    zip_path = os.path.join(tmp_dir, "resumes.zip")

    with zipfile.ZipFile(zip_path, "w") as zipf:
        for jd in jds:
            tailored = db.query(TailoredResume).filter(TailoredResume.jd_id == jd.id).first()
            if not tailored:
                continue
            title = jd.title if jd.title else f"resume_{jd.id}"
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()
            docx_path = os.path.join(tmp_dir, f"{safe_title}.docx")
            tailored_content = json.loads(tailored.content)
            build_docx(tailored_content, docx_path)
            zipf.write(docx_path, f"{safe_title}.docx")

    return FileResponse(
        path=zip_path,
        filename="tailored_resumes.zip",
        media_type="application/zip"
    )